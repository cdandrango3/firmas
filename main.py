import docx
from docx.shared import Inches
dias=["Lunes.docx","Martes.docx","Miercoles.docx","Jueves.docx","Viernes.docx"]
def docume(url):
  doc=docx.Document("firmar docu/{0}".format(url))
  table=doc.tables
  for i in range (2,12):
    p=table[0].rows[i].cells[6].add_paragraph()
    d=p.add_run()
    d.add_picture("firmasim/{0}.png".format(i-2),width=Inches(1), height=Inches(0.6))

  doc.save("documentos corregidos " + url)
for c in dias:
  docume(c)
